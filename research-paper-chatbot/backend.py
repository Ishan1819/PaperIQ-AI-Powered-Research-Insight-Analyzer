# import re
# import fitz  # PyMuPDF
# import nltk
# import math
# import numpy as np

# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity

# from nltk.corpus import stopwords, wordnet
# from nltk.stem import WordNetLemmatizer
# from nltk.tokenize import sent_tokenize, word_tokenize

# from io import BytesIO
# from docx import Document


# # Download required NLTK data
# try:
#     nltk.data.find('tokenizers/punkt')
#     nltk.data.find('taggers/averaged_perceptron_tagger')
#     nltk.data.find('corpora/wordnet')
#     nltk.data.find('corpora/stopwords')
# except LookupError:
#     nltk.download('punkt', quiet=True)
#     nltk.download('averaged_perceptron_tagger', quiet=True)
#     nltk.download('wordnet', quiet=True)
#     nltk.download('stopwords', quiet=True)
#     nltk.download('omw-1.4', quiet=True)

# stop_words = set(stopwords.words("english"))
# lemmatizer = WordNetLemmatizer()


# def get_wordnet_pos(tag: str):
#     if tag.startswith("J"):
#         return wordnet.ADJ
#     elif tag.startswith("V"):
#         return wordnet.VERB
#     elif tag.startswith("N"):
#         return wordnet.NOUN
#     elif tag.startswith("R"):
#         return wordnet.ADV
#     else:
#         return wordnet.NOUN


# def preprocess_sentence(sentence: str) -> str:
#     sentence = ' '.join([w for w in sentence.split() 
#                         if not (w.replace('.', '').isdigit() or w.strip('()').isdigit())])
#     words = word_tokenize(sentence.lower())
#     pos_tags = nltk.pos_tag(words)
#     lemmatized = [lemmatizer.lemmatize(word, get_wordnet_pos(tag)) 
#                   for word, tag in pos_tags if word.isalpha() and word not in stop_words]
#     return " ".join(lemmatized)


# def clean_section_content(text: str) -> str:
#     lines = []
#     for line in text.split('\n'):
#         line = line.strip()
#         line = ' '.join([w for w in line.split() 
#                         if not (w.replace('.', '').isdigit() or w.strip('()').isdigit())])
#         if line:
#             lines.append(line)
#     return ' '.join(lines)


# def extract_text_from_pdf(pdf_path) -> str:
#     doc = fitz.open(pdf_path)
#     full_text = []
#     for page in doc:
#         text = page.get_text()
#         text = text.replace("\n\n", "[PARAGRAPH]").replace("\n", " ").replace("[PARAGRAPH]", "\n\n")
#         full_text.append(text)
#     doc.close()
#     return " ".join(full_text)


# def extract_text_from_docx(docx_path) -> str:
#     doc = Document(docx_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return "\n".join(full_text)


# def split_into_paragraphs(clean_text: str):
#     chunk_size = 200
#     paragraphs = []
#     sentences = sent_tokenize(clean_text)
#     sentences = sentences[1:] if len(sentences) > 1 else []

#     current_para = []
#     current_length = 0

#     for sentence in sentences:
#         current_para.append(sentence)
#         current_length += len(sentence) + 1

#         if current_length >= chunk_size:
#             chunk_text = ' '.join(current_para)
#             paragraphs.append(chunk_text)
#             current_para = []
#             current_length = 0

#     if current_para:
#         paragraphs.append(' '.join(current_para))

#     return paragraphs


# def extract_sections(text: str) -> dict:
#     sections = {}
#     current_section = "Main"
#     current_content = []
#     lines = text.split("\n")

#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue
#         if line.isupper() or line.endswith(":"):
#             if current_content:
#                 section_text = " ".join(current_content)
#                 clean_text = clean_section_content(section_text)
#                 paragraphs = split_into_paragraphs(clean_text)
#                 sections[current_section] = paragraphs
#             current_section = clean_section_content(line.rstrip(":"))
#             current_content = []
#         else:
#             current_content.append(line)

#     if current_content:
#         section_text = " ".join(current_content)
#         clean_text = clean_section_content(section_text)
#         paragraphs = split_into_paragraphs(clean_text)
#         sections[current_section] = paragraphs

#     return sections


# def extract_topic_from_query(query: str) -> str:
#     query_lower = query.lower()
#     query_cleaned = re.sub(r'\b(what is|what are|define|tell me about|explain|describe|give me|show me)\b', 
#                           '', query_lower, flags=re.IGNORECASE)
#     query_cleaned = re.sub(r'\b(the|a|an|of|in|on|at|to|for)\b', '', query_cleaned)
#     query_cleaned = query_cleaned.strip('? .')
#     return query_cleaned.strip()


# def find_topic_in_document(topic: str, full_text: str, sections: dict) -> tuple:
#     topic_normalized = topic.strip().lower()
#     all_paragraphs = [para for paras in sections.values() for para in paras]

#     pattern1 = re.compile(rf'\b{re.escape(topic_normalized)}[:\s].*?key features', re.IGNORECASE | re.DOTALL)
#     pattern2 = re.compile(rf'(?:^|\.\s+){re.escape(topic_normalized)}\s*:', re.IGNORECASE)
#     pattern3 = re.compile(rf'\b{re.escape(topic_normalized)}\s+(is|are|was|were)\b', re.IGNORECASE)
#     pattern4 = re.compile(rf'(?:^|\n){re.escape(topic_normalized)}\b', re.IGNORECASE)
#     pattern5 = re.compile(rf'\b{re.escape(topic_normalized)}\b', re.IGNORECASE)
    
#     matches = []
    
#     for para_idx, para in enumerate(all_paragraphs):
#         score = 0
#         match_type = None
        
#         if pattern1.search(para):
#             score = 100
#             match_type = "complete_section"
#         elif pattern2.search(para):
#             score = 80
#             match_type = "definition"
#         elif pattern3.search(para):
#             score = 60
#             match_type = "is_pattern"
#         elif pattern4.search(para):
#             score = 40
#             match_type = "sentence_start"
#         elif pattern5.search(para):
#             count = len(pattern5.findall(para))
#             score = 20 * count
#             match_type = "general"
        
#         if score > 0:
#             matches.append((para, match_type, score, para_idx))
    
#     if matches:
#         matches.sort(key=lambda x: x[2], reverse=True)
#         best_match, best_type, best_score, idx = matches[0]
#         return True, best_match, best_type
    
#     return False, None, None


# def generate_natural_response(query: str, extracted_content: str, topic: str) -> str:
#     query_lower = query.lower()

#     if any(word in query_lower for word in ['what is', 'what are', 'define', 'meaning']):
#         if topic:
#             return f"{topic.title()} is {extracted_content}"
#         return extracted_content
    
#     elif any(word in query_lower for word in ['key features', 'features', 'characteristics']):
#         if 'key features' in extracted_content.lower():
#             return extracted_content
#         return f"The key features of {topic.title()} include: {extracted_content}"
    
#     elif any(word in query_lower for word in ['how does', 'how do', 'working', 'works']):
#         return f"{topic.title()} works as follows: {extracted_content}"
    
#     elif 'why' in query_lower:
#         return f"Regarding why {topic}: {extracted_content}"
    
#     elif any(word in query_lower for word in ['list', 'enumerate', 'what are']):
#         return f"Here are the details about {topic}: {extracted_content}"
#     else:
#         return f"Regarding {topic}: {extracted_content}" if topic else extracted_content


# def format_response(query: str, paragraph: str, match_type: str, topic: str) -> str:
#     query_lower = query.lower()
    
#     is_definition_query = any(word in query_lower for word in ['what is', 'what are', 'define', 'definition', 'meaning'])
#     is_features_query = any(word in query_lower for word in ['key features', 'features', 'characteristics', 'properties'])
#     is_how_query = any(word in query_lower for word in ['how', 'working', 'works', 'process', 'mechanism'])
    
#     sentences = sent_tokenize(paragraph)

#     if is_definition_query:
#         for sentence in sentences:
#             if re.search(rf'\b{re.escape(topic)}\s+(is|are)\b', sentence, re.IGNORECASE):
#                 idx = sentences.index(sentence)
#                 return ' '.join(sentences[idx:min(idx+2, len(sentences))])
#         return ' '.join(sentences[:3])
    
#     elif is_features_query:
#         features_match = re.search(r'key features.*?:', paragraph, re.IGNORECASE)
        
#         if features_match:
#             start_pos = features_match.end()
#             features_text = paragraph[start_pos:].strip()
#             feature_sentences = sent_tokenize(features_text)
            
#             actual_features = []
#             for sent in feature_sentences:
#                 if re.match(r'^[A-Z][a-z]+\s*:', sent) or sent.isupper():
#                     break
#                 actual_features.append(sent)
            
#             if actual_features:
#                 return "Key features: " + ' '.join(actual_features)
        
#         feature_sentences = []
#         for sent in sentences:
#             if re.search(rf'\b{re.escape(topic)}\s+(is|are)\b', sent, re.IGNORECASE):
#                 continue
#             if len(sent) < 200 or any(word in sent.lower() for word in ['support', 'provide', 'include', 'enable', 'feature']):
#                 feature_sentences.append(sent)
        
#         if feature_sentences:
#             return ' '.join(feature_sentences[:6])
#         return ' '.join(sentences[1:5]) if len(sentences) > 1 else ' '.join(sentences)
    
#     elif is_how_query:
#         return ' '.join(sentences[:4])
    
#     else:
#         if match_type == "complete_section":
#             return ' '.join(sentences[:5])
        
#         if match_type == "is_pattern":
#             for sentence in sentences:
#                 if re.search(rf'\b{re.escape(topic)}\s+(is|are|was|were)\b', sentence, re.IGNORECASE):
#                     idx = sentences.index(sentence)
#                     return ' '.join(sentences[idx:min(idx+3, len(sentences))])
        
#         return ' '.join(sentences[:3])


# def simple_summarize(text, summary_ratio=0.3):
#     sentences = sent_tokenize(text)
#     cleaned_sentences = []
#     for s in sentences:
#         words = word_tokenize(s.lower())
#         cleaned_sent = [w for w in words if w.isalpha() and w not in stop_words]
#         cleaned_sentences.append(cleaned_sent)
    
#     vocab = set([w for sent in cleaned_sentences for w in sent])
#     vocab = list(vocab)
    
#     idf = {}
#     total_sentences = len(sentences)
#     for word in vocab:
#         count = sum(1 for sent in cleaned_sentences if word in sent)
#         if count > 0:
#             idf[word] = math.log(total_sentences / count)
    
#     sentence_scores = []
#     for sent, words in zip(sentences, cleaned_sentences):
#         if len(words) > 0:
#             tfidf_score = sum((words.count(word)/len(words))*idf.get(word, 0) for word in words)
#             sentence_scores.append((sent, tfidf_score))
#         else:
#             sentence_scores.append((sent, 0))
    
#     num_sentences = max(1, int(len(sentences) * summary_ratio))
#     ranked = sorted(sentence_scores, key=lambda x: x[1], reverse=True)
#     selected = [item[0] for item in ranked[:num_sentences]]
    
#     return " ".join(selected)


# def retrieve_paragraph_answer(user_query: str, sections: dict, full_text: str, threshold: float=0.15) -> str:
#     topic = extract_topic_from_query(user_query)
    
#     if topic:
#         found, paragraph, match_type = find_topic_in_document(topic, full_text, sections)
        
#         if found:
#             extracted_content = format_response(user_query, paragraph, match_type, topic)
#             natural_response = generate_natural_response(user_query, extracted_content, topic)
#             return natural_response
    
#     all_paragraphs = [para for paras in sections.values() for para in paras]
#     all_sentences = []
#     sentence_to_para_idx = {}
    
#     for para_idx, paragraph in enumerate(all_paragraphs):
#         sentences = sent_tokenize(paragraph)
#         for sentence in sentences:
#             sentence_to_para_idx[len(all_sentences)] = para_idx
#             all_sentences.append(sentence)
    
#     if not all_sentences:
#         return "No content found in the document."

#     processed_sentences = [preprocess_sentence(s) for s in all_sentences]
#     processed_query = preprocess_sentence(user_query)
    
#     if not processed_query.strip():
#         return "Please provide a more specific query."

#     vectorizer = TfidfVectorizer()
#     vectors = vectorizer.fit_transform([processed_query] + processed_sentences)
#     similarity_scores = cosine_similarity(vectors[0:1], vectors[1:]).flatten()
#     best_sentence_idx = similarity_scores.argmax()
#     best_score = similarity_scores[best_sentence_idx]

#     if best_score > threshold:
#         best_para_idx = sentence_to_para_idx[best_sentence_idx]
#         retrieved = all_paragraphs[best_para_idx]

#         extracted_content = format_response(user_query, retrieved, "general", topic)
#         natural_response = generate_natural_response(user_query, extracted_content, topic)
#         return natural_response
#     else:
#         return "No relevant information found. Try rephrasing your question."








































import re
import fitz  # PyMuPDF
import nltk
import math
import numpy as np

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from nltk.corpus import stopwords, wordnet
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import sent_tokenize, word_tokenize

from io import BytesIO
from docx import Document


# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('taggers/averaged_perceptron_tagger')
    nltk.data.find('corpora/wordnet')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
    nltk.download('wordnet', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('omw-1.4', quiet=True)

stop_words = set(stopwords.words("english"))
lemmatizer = WordNetLemmatizer()


def get_wordnet_pos(tag: str):
    if tag.startswith("J"):
        return wordnet.ADJ
    elif tag.startswith("V"):
        return wordnet.VERB
    elif tag.startswith("N"):
        return wordnet.NOUN
    elif tag.startswith("R"):
        return wordnet.ADV
    else:
        return wordnet.NOUN


def preprocess_sentence(sentence: str) -> str:
    sentence = ' '.join([w for w in sentence.split() 
                        if not (w.replace('.', '').isdigit() or w.strip('()').isdigit())])
    words = word_tokenize(sentence.lower())
    pos_tags = nltk.pos_tag(words)
    lemmatized = [lemmatizer.lemmatize(word, get_wordnet_pos(tag)) 
                  for word, tag in pos_tags if word.isalpha() and word not in stop_words]
    return " ".join(lemmatized)


def clean_section_content(text: str) -> str:
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        line = ' '.join([w for w in line.split() 
                        if not (w.replace('.', '').isdigit() or w.strip('()').isdigit())])
        if line:
            lines.append(line)
    return ' '.join(lines)



def extract_text_from_docx(docx_path) -> tuple:
    """Extract text from DOCX - page numbers are approximate based on paragraphs"""
    doc = Document(docx_path)
    pages_data = []
    
    # Approximate pages (every 10 paragraphs = 1 page, adjust as needed)
    paragraphs_per_page = 10
    current_page = 1
    page_content = []
    
    for idx, para in enumerate(doc.paragraphs):
        page_content.append(para.text)
        
        if (idx + 1) % paragraphs_per_page == 0:
            pages_data.append({
                "page": current_page,
                "content": "\n".join(page_content)
            })
            current_page += 1
            page_content = []
    
    # Add remaining content
    if page_content:
        pages_data.append({
            "page": current_page,
            "content": "\n".join(page_content)
        })
    
    full_text = "\n".join([p["content"] for p in pages_data])
    return full_text, pages_data


def split_into_paragraphs(clean_text: str, page_num: int):
    """Split text into paragraphs with metadata"""
    chunk_size = 200
    paragraphs = []
    sentences = sent_tokenize(clean_text)
    sentences = sentences[1:] if len(sentences) > 1 else []

    current_para = []
    current_length = 0

    for sentence in sentences:
        current_para.append(sentence)
        current_length += len(sentence) + 1

        if current_length >= chunk_size:
            chunk_text = ' '.join(current_para)
            paragraphs.append({
                "text": chunk_text,
                "page": page_num
            })
            current_para = []
            current_length = 0

    if current_para:
        paragraphs.append({
            "text": ' '.join(current_para),
            "page": page_num
        })

    return paragraphs


def extract_text_from_pdf(pdf_path) -> tuple:
    """Extract text from PDF with precise page-to-text mapping"""
    doc = fitz.open(pdf_path)
    pages_data = []
    
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        # Keep newlines for better structure preservation
        text = text.strip()
        if text:  # Only add non-empty pages
            pages_data.append({
                "page": page_num,
                "content": text
            })
    
    doc.close()
    full_text = "\n\n".join([f"[PAGE_{p['page']}]\n{p['content']}" for p in pages_data])
    return full_text, pages_data

def remove_section_headers(text: str) -> str:
    """Remove section headers that end with colon from the beginning or middle of text"""
    # Pattern to match headers like "Conclusion:", "Introduction:", "Alpha Code:", etc.
    # Matches capitalized words followed by colon at start or after period/newline
    text = re.sub(r'(?:^|\.\s+)([A-Z][A-Za-z\s]+):\s*', '', text)
    text = re.sub(r'^([A-Z][A-Z\s]+):\s*', '', text)  # All caps headers
    return text.strip()


def extract_sections_precise(pages_data: list) -> dict:
    """Extract sections with precise page tracking and header filtering"""
    sections = {}
    
    # First, create chunks directly from each page
    all_chunks = []
    
    for page_data in pages_data:
        page_num = page_data["page"]
        text = page_data["content"]
        
        # Split into lines and filter out headers
        lines = text.split('\n')
        content_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip section headers
            if line.isupper() or (line.endswith(':') and len(line.split()) <= 5):
                continue
            if line:
                content_lines.append(line)
        
        # Join filtered content
        filtered_text = ' '.join(content_lines)
        
        # Clean the text
        cleaned = clean_section_content(filtered_text)
        
        # Split into sentences
        sentences = sent_tokenize(cleaned)
        
        # Create chunks from this page's sentences
        chunk_size = 200
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            current_chunk.append(sentence)
            current_length += len(sentence) + 1
            
            if current_length >= chunk_size:
                chunk_text = ' '.join(current_chunk)
                # Final cleanup to remove any remaining headers
                chunk_text = remove_section_headers(chunk_text)
                
                all_chunks.append({
                    "text": chunk_text,
                    "page": page_num
                })
                current_chunk = []
                current_length = 0
        
        # Add remaining sentences from this page
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunk_text = remove_section_headers(chunk_text)
            
            all_chunks.append({
                "text": chunk_text,
                "page": page_num
            })
    
    # Organize chunks into sections (or just use "Main" section)
    sections["Main"] = all_chunks
    
    return sections


def extract_topic_from_query(query: str) -> str:
    query_lower = query.lower()
    query_cleaned = re.sub(r'\b(what is|what are|define|tell me about|explain|describe|give me|show me)\b', 
                          '', query_lower, flags=re.IGNORECASE)
    query_cleaned = re.sub(r'\b(the|a|an|of|in|on|at|to|for)\b', '', query_cleaned)
    query_cleaned = query_cleaned.strip('? .')
    return query_cleaned.strip()


def find_topic_in_document(topic: str, full_text: str, sections: dict) -> tuple:
    """Find topic in document and return with metadata"""
    topic_normalized = topic.strip().lower()
    all_paragraphs = [para for paras in sections.values() for para in paras]

    pattern1 = re.compile(rf'\b{re.escape(topic_normalized)}[:\s].*?key features', re.IGNORECASE | re.DOTALL)
    pattern2 = re.compile(rf'(?:^|\.\s+){re.escape(topic_normalized)}\s*:', re.IGNORECASE)
    pattern3 = re.compile(rf'\b{re.escape(topic_normalized)}\s+(is|are|was|were)\b', re.IGNORECASE)
    pattern4 = re.compile(rf'(?:^|\n){re.escape(topic_normalized)}\b', re.IGNORECASE)
    pattern5 = re.compile(rf'\b{re.escape(topic_normalized)}\b', re.IGNORECASE)
    
    matches = []
    
    for para_idx, para_dict in enumerate(all_paragraphs):
        para = para_dict["text"]
        page_num = para_dict["page"]
        score = 0
        match_type = None
        
        if pattern1.search(para):
            score = 100
            match_type = "complete_section"
        elif pattern2.search(para):
            score = 80
            match_type = "definition"
        elif pattern3.search(para):
            score = 60
            match_type = "is_pattern"
        elif pattern4.search(para):
            score = 40
            match_type = "sentence_start"
        elif pattern5.search(para):
            count = len(pattern5.findall(para))
            score = 20 * count
            match_type = "general"
        
        if score > 0:
            matches.append((para, match_type, score, para_idx, page_num))
    
    if matches:
        matches.sort(key=lambda x: x[2], reverse=True)
        best_match, best_type, best_score, idx, page_num = matches[0]
        return True, best_match, best_type, page_num
    
    return False, None, None, None


def generate_natural_response(query: str, extracted_content: str, topic: str, page_num: int = None) -> str:
    """Generate natural response with page citation"""
    query_lower = query.lower()
    
    # Build the response based on query type
    if any(word in query_lower for word in ['what is', 'what are', 'define', 'meaning']):
        if topic:
            response = f"{topic.title()} is {extracted_content}"
        else:
            response = extracted_content
    
    elif any(word in query_lower for word in ['key features', 'features', 'characteristics']):
        if 'key features' in extracted_content.lower():
            response = extracted_content
        else:
            response = f"The key features of {topic.title()} include: {extracted_content}"
    
    elif any(word in query_lower for word in ['how does', 'how do', 'working', 'works']):
        response = f"{topic.title()} works as follows: {extracted_content}"
    
    elif 'why' in query_lower:
        response = f"Regarding why {topic}: {extracted_content}"
    
    elif any(word in query_lower for word in ['list', 'enumerate', 'what are']):
        response = f"Here are the details about {topic}: {extracted_content}"
    else:
        response = f"Regarding {topic}: {extracted_content}" if topic else extracted_content
    
    # Add page citation if available
    if page_num:
        response += f"\n\n[Source: Page {page_num}]"
    
    return response


def format_response(query: str, paragraph: str, match_type: str, topic: str) -> str:
    query_lower = query.lower()
    
    is_definition_query = any(word in query_lower for word in ['what is', 'what are', 'define', 'definition', 'meaning'])
    is_features_query = any(word in query_lower for word in ['key features', 'features', 'characteristics', 'properties'])
    is_how_query = any(word in query_lower for word in ['how', 'working', 'works', 'process', 'mechanism'])
    
    sentences = sent_tokenize(paragraph)

    if is_definition_query:
        for sentence in sentences:
            if re.search(rf'\b{re.escape(topic)}\s+(is|are)\b', sentence, re.IGNORECASE):
                idx = sentences.index(sentence)
                return ' '.join(sentences[idx:min(idx+2, len(sentences))])
        return ' '.join(sentences[:3])
    
    elif is_features_query:
        features_match = re.search(r'key features.*?:', paragraph, re.IGNORECASE)
        
        if features_match:
            start_pos = features_match.end()
            features_text = paragraph[start_pos:].strip()
            feature_sentences = sent_tokenize(features_text)
            
            actual_features = []
            for sent in feature_sentences:
                if re.match(r'^[A-Z][a-z]+\s*:', sent) or sent.isupper():
                    break
                actual_features.append(sent)
            
            if actual_features:
                return "Key features: " + ' '.join(actual_features)
        
        feature_sentences = []
        for sent in sentences:
            if re.search(rf'\b{re.escape(topic)}\s+(is|are)\b', sent, re.IGNORECASE):
                continue
            if len(sent) < 200 or any(word in sent.lower() for word in ['support', 'provide', 'include', 'enable', 'feature']):
                feature_sentences.append(sent)
        
        if feature_sentences:
            return ' '.join(feature_sentences[:6])
        return ' '.join(sentences[1:5]) if len(sentences) > 1 else ' '.join(sentences)
    
    elif is_how_query:
        return ' '.join(sentences[:4])
    
    else:
        if match_type == "complete_section":
            return ' '.join(sentences[:5])
        
        if match_type == "is_pattern":
            for sentence in sentences:
                if re.search(rf'\b{re.escape(topic)}\s+(is|are|was|were)\b', sentence, re.IGNORECASE):
                    idx = sentences.index(sentence)
                    return ' '.join(sentences[idx:min(idx+3, len(sentences))])
        
        return ' '.join(sentences[:3])


def simple_summarize(text, summary_ratio=0.3):
    sentences = sent_tokenize(text)
    cleaned_sentences = []
    for s in sentences:
        words = word_tokenize(s.lower())
        cleaned_sent = [w for w in words if w.isalpha() and w not in stop_words]
        cleaned_sentences.append(cleaned_sent)
    
    vocab = set([w for sent in cleaned_sentences for w in sent])
    vocab = list(vocab)
    
    idf = {}
    total_sentences = len(sentences)
    for word in vocab:
        count = sum(1 for sent in cleaned_sentences if word in sent)
        if count > 0:
            idf[word] = math.log(total_sentences / count)
    
    sentence_scores = []
    for sent, words in zip(sentences, cleaned_sentences):
        if len(words) > 0:
            tfidf_score = sum((words.count(word)/len(words))*idf.get(word, 0) for word in words)
            sentence_scores.append((sent, tfidf_score))
        else:
            sentence_scores.append((sent, 0))
    
    num_sentences = max(1, int(len(sentences) * summary_ratio))
    ranked = sorted(sentence_scores, key=lambda x: x[1], reverse=True)
    selected = [item[0] for item in ranked[:num_sentences]]
    
    return " ".join(selected)


def retrieve_paragraph_answer(user_query: str, sections: dict, full_text: str, threshold: float=0.15) -> str:
    """Retrieve answer with page citation"""
    topic = extract_topic_from_query(user_query)
    
    if topic:
        found, paragraph, match_type, page_num = find_topic_in_document(topic, full_text, sections)
        
        if found:
            extracted_content = format_response(user_query, paragraph, match_type, topic)
            natural_response = generate_natural_response(user_query, extracted_content, topic, page_num)
            return natural_response
    
    all_paragraphs = [para for paras in sections.values() for para in paras]
    all_sentences = []
    sentence_to_para_idx = {}
    
    for para_idx, para_dict in enumerate(all_paragraphs):
        paragraph = para_dict["text"]
        sentences = sent_tokenize(paragraph)
        for sentence in sentences:
            sentence_to_para_idx[len(all_sentences)] = para_idx
            all_sentences.append(sentence)
    
    if not all_sentences:
        return "No content found in the document."

    processed_sentences = [preprocess_sentence(s) for s in all_sentences]
    processed_query = preprocess_sentence(user_query)
    
    if not processed_query.strip():
        return "Please provide a more specific query."

    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([processed_query] + processed_sentences)
    similarity_scores = cosine_similarity(vectors[0:1], vectors[1:]).flatten()
    best_sentence_idx = similarity_scores.argmax()
    best_score = similarity_scores[best_sentence_idx]

    if best_score > threshold:
        best_para_idx = sentence_to_para_idx[best_sentence_idx]
        retrieved_dict = all_paragraphs[best_para_idx]
        retrieved = retrieved_dict["text"]
        page_num = retrieved_dict["page"]

        extracted_content = format_response(user_query, retrieved, "general", topic)
        natural_response = generate_natural_response(user_query, extracted_content, topic, page_num)
        return natural_response
    else:
        return "No relevant information found. Try rephrasing your question."
